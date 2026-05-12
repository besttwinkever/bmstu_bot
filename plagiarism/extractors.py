"""Извлечение сырого текста из файлов работ.

Поддерживаем форматы, которые реально встречаются в учебных работах:
`.txt`, `.docx`, `.pdf`. Неизвестный формат → `UnsupportedFormat`, вызов
сам решит: пропустить работу или отметить отчёт как несостоявшийся.
"""
from __future__ import annotations

import io
import os
from typing import Callable, Dict


class UnsupportedFormat(Exception):
    """Расширение файла не поддерживается для извлечения текста."""


def _from_txt(data: bytes) -> str:
    for encoding in ('utf-8', 'cp1251', 'latin-1'):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode('utf-8', errors='ignore')


def _from_docx(data: bytes) -> str:
    from docx import Document  # python-docx
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise UnsupportedFormat(f'Файл .docx повреждён или не читается: {exc}')
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return '\n'.join(parts)


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError
    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as exc:
        raise UnsupportedFormat(f'Файл .pdf повреждён или не читается: {exc}')
    if reader.is_encrypted:
        raise UnsupportedFormat('Файл .pdf зашифрован — текст недоступен для проверки.')
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or '')
        except Exception:
            # Один битый Page не должен валить весь файл.
            continue
    return '\n'.join(pages)


_EXTRACTORS: Dict[str, Callable[[bytes], str]] = {
    '.txt': _from_txt,
    '.docx': _from_docx,
    '.pdf': _from_pdf,
}


def extract_text(file_name: str, data: bytes) -> str:
    ext = os.path.splitext(file_name)[1].lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        supported = ', '.join(sorted(_EXTRACTORS.keys())) or 'нет'
        raise UnsupportedFormat(
            f'Формат «{ext or "—"}» не поддерживается. '
            f'Поддерживаются: {supported}.'
        )
    return extractor(data)


def supported_extensions() -> list[str]:
    return [ext.lstrip('.') for ext in _EXTRACTORS]
